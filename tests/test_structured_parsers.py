from __future__ import annotations

import ast
import inspect
import math
import os
import subprocess
from pathlib import Path
from types import SimpleNamespace

import pytest

from KBzhy.app.core.document_models import DocumentElement, ParsedDocument
from KBzhy.app.core.parser import (
    Document,
    DocumentParseError,
    DocumentParser,
    ParsedArtifactError,
)


def _markdown_fixture(tmp_path: Path) -> Path:
    path = tmp_path / "guide.md"
    path.write_text("# Guide\n\nInstall the package.", encoding="utf-8")
    return path


def test_parse_structured_returns_parsed_document_and_keeps_legacy_parse(tmp_path):
    source = _markdown_fixture(tmp_path)
    parser = DocumentParser()

    legacy = parser.parse(source)
    parsed = parser.parse_structured(
        source,
        document_id="doc1",
        version=2,
        kb_id="kb1",
    )

    assert isinstance(legacy, list)
    assert all(isinstance(document, Document) for document in legacy)
    assert parsed.document_id == "doc1"
    assert parsed.version == 2
    assert parsed.title == "guide"
    assert parsed.language == "und"
    assert parsed.metadata["kb_id"] == "kb1"
    assert parsed.elements
    assert [element.element_type for element in parsed.elements] == ["heading", "paragraph"]
    assert parsed.elements[1].section_path == ("Guide",)
    assert legacy[0].content == "# Guide\n\nInstall the package."


def test_markdown_structured_parser_preserves_sections_code_lists_and_tables(tmp_path):
    source = tmp_path / "api.md"
    source.write_text(
        "# API\n\n## 创建用户\n\n- 校验参数\n- 写入数据\n\n"
        "```python\ncreate_user()\n```\n\n"
        "| 字段 | 类型 |\n| --- | --- |\n| name | string |\n",
        encoding="utf-8",
    )

    parsed = DocumentParser().parse_structured(
        source, document_id="doc-md", version=1, kb_id="kb1"
    )

    assert [element.element_type for element in parsed.elements] == [
        "heading",
        "heading",
        "list",
        "code",
        "table",
    ]
    code = parsed.elements[3]
    assert code.text == "```python\ncreate_user()\n```"
    assert code.section_path == ("API", "创建用户")
    assert parsed.elements[4].text == (
        "| 字段 | 类型 |\n| --- | --- |\n| name | string |"
    )
    assert [element.order for element in parsed.elements] == list(range(5))


def test_markdown_fence_contents_are_not_reparsed_as_markdown(tmp_path):
    source = tmp_path / "sample.md"
    source.write_text("# Outside\n\n```md\n# Inside\n- item\n```", encoding="utf-8")

    parsed = DocumentParser().parse_structured(source, document_id="doc-md", version=1)

    assert [element.element_type for element in parsed.elements] == ["heading", "code"]
    assert parsed.elements[1].section_path == ("Outside",)


def test_markdown_heading_stack_discards_deeper_sibling_context(tmp_path):
    source = tmp_path / "sections.md"
    source.write_text("### Deep\n\n## Parent\n\nbody", encoding="utf-8")

    parsed = DocumentParser().parse_structured(source, document_id="doc-md", version=1)

    assert parsed.elements[1].section_path == ("Parent",)
    assert parsed.elements[2].section_path == ("Parent",)


def test_markdown_table_stops_before_markdown_block_with_pipe(tmp_path):
    source = tmp_path / "table.md"
    source.write_text(
        "| A | B |\n| --- | --- |\n| 1 | 2 |\n# Next | Section\n\nbody",
        encoding="utf-8",
    )

    parsed = DocumentParser().parse_structured(source, document_id="doc-md", version=1)

    assert [element.element_type for element in parsed.elements] == [
        "table",
        "heading",
        "paragraph",
    ]
    assert parsed.elements[0].text.endswith("| 1 | 2 |")
    assert parsed.elements[1].text == "Next | Section"


def test_markdown_table_rejects_inconsistent_rows(tmp_path):
    source = tmp_path / "table.md"
    source.write_text(
        "| A | B |\n| --- | --- |\n| valid | row |\n| incomplete |\n",
        encoding="utf-8",
    )

    parsed = DocumentParser().parse_structured(source, document_id="doc-md", version=1)

    assert parsed.elements[0].element_type == "table"
    assert "incomplete" not in parsed.elements[0].text
    assert parsed.elements[1].element_type == "paragraph"
    assert parsed.elements[1].text == "| incomplete |"


def test_markdown_list_keeps_indented_continuation_in_same_element(tmp_path):
    source = tmp_path / "list.md"
    source.write_text(
        "- first item\n  continuation line\n  still first\n- second item",
        encoding="utf-8",
    )

    parsed = DocumentParser().parse_structured(source, document_id="doc-md", version=1)

    assert len(parsed.elements) == 1
    assert parsed.elements[0].element_type == "list"
    assert parsed.elements[0].text == (
        "- first item\n  continuation line\n  still first\n- second item"
    )


@pytest.mark.parametrize(
    "heading,expected",
    [
        ("# C#", "C#"),
        ("# C# ###", "C#"),
        ("# 标题###", "标题###"),
        ("# 标题 ###", "标题"),
    ],
)
def test_markdown_atx_closing_sequence_requires_preceding_space(
    tmp_path, heading, expected
):
    source = tmp_path / "heading.md"
    source.write_text(heading, encoding="utf-8")

    parsed = DocumentParser().parse_structured(source, document_id="doc-md", version=1)

    assert parsed.elements[0].text == expected


def test_txt_parser_uses_conservative_headings_and_keeps_plain_short_lines_as_paragraphs(
    tmp_path,
):
    source = tmp_path / "notes.txt"
    source.write_text(
        "这是短句\n\n第一章 总则\n\n第一段正文。\n\n第二段正文。",
        encoding="utf-8",
    )

    parsed = DocumentParser().parse_structured(source, document_id="doc-txt", version=1)

    assert [element.element_type for element in parsed.elements] == [
        "paragraph",
        "heading",
        "paragraph",
        "paragraph",
    ]
    assert parsed.elements[0].section_path == ()
    assert parsed.elements[2].section_path == ("第一章 总则",)


def test_structured_text_parser_rejects_empty_content(tmp_path):
    source = tmp_path / "empty.txt"
    source.write_text(" \n\n\t", encoding="utf-8")

    with pytest.raises(DocumentParseError, match="未提取到可索引文字"):
        DocumentParser().parse_structured(source, document_id="doc-empty", version=1)


@pytest.mark.parametrize(
    "payload",
    [b"abc\x01def", b"\x01\x02\x03binary", b"abc\xef\xbf\xbddef"],
)
def test_text_parser_rejects_forbidden_control_characters(tmp_path, payload):
    source = tmp_path / "binary.txt"
    source.write_bytes(payload)

    with pytest.raises(DocumentParseError, match="文本质量"):
        DocumentParser().parse_structured(source, document_id="doc-text", version=1)


def test_text_parser_accepts_normal_chinese_english_and_whitespace(tmp_path):
    source = tmp_path / "normal.txt"
    source.write_text("中文 English\tvalue\n\n下一段", encoding="utf-8")

    parsed = DocumentParser().parse_structured(source, document_id="doc-text", version=1)

    assert [element.text for element in parsed.elements] == [
        "中文 English\tvalue",
        "下一段",
    ]


def test_text_quality_validation_does_not_materialize_character_collections():
    from KBzhy.app.core.parsers.text_parser import _validate_text_quality

    function = ast.parse(inspect.getsource(_validate_text_quality))

    assert not any(
        isinstance(node, (ast.ListComp, ast.SetComp, ast.DictComp))
        for node in ast.walk(function)
    )
    assert not any(
        isinstance(node, ast.Call)
        and isinstance(node.func, ast.Name)
        and node.func.id in {"list", "tuple", "set", "dict"}
        for node in ast.walk(function)
    )


def test_word_parser_preserves_body_order_heading_path_lists_and_header_context(tmp_path):
    from docx import Document as WordDocument

    source = tmp_path / "guide.docx"
    document = WordDocument()
    document.add_heading("指南", level=1)
    document.add_paragraph("开场正文")
    table = document.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "字段"
    table.rows[0].cells[1].text = "类型"
    table.rows[1].cells[0].text = "name"
    table.rows[1].cells[1].text = "string"
    document.add_heading("步骤", level=2)
    document.add_paragraph("创建用户", style="List Bullet")
    document.save(source)

    parsed = DocumentParser().parse_structured(source, document_id="doc-word", version=1)

    assert [element.element_type for element in parsed.elements] == [
        "heading",
        "paragraph",
        "table",
        "heading",
        "list",
    ]
    assert parsed.elements[2].text == (
        "| 字段 | 类型 |\n| --- | --- |\n| name | string |"
    )
    assert parsed.elements[2].section_path == ("指南",)
    assert parsed.elements[4].section_path == ("指南", "步骤")


def test_word_heading_stack_discards_deeper_sibling_context(tmp_path):
    from docx import Document as WordDocument

    source = tmp_path / "sections.docx"
    document = WordDocument()
    document.add_heading("Deep", level=3)
    document.add_heading("Parent", level=2)
    document.add_paragraph("body")
    document.save(source)

    parsed = DocumentParser().parse_structured(source, document_id="doc-word", version=1)

    assert parsed.elements[1].section_path == ("Parent",)
    assert parsed.elements[2].section_path == ("Parent",)


def test_word_parser_rejects_empty_document(tmp_path):
    from docx import Document as WordDocument

    source = tmp_path / "empty.docx"
    WordDocument().save(source)

    with pytest.raises(DocumentParseError, match="未提取到可索引文字"):
        DocumentParser().parse_structured(source, document_id="doc-empty", version=1)


def test_word_parser_normalizes_horizontal_and_vertical_merged_cells(tmp_path):
    from docx import Document as WordDocument

    source = tmp_path / "merged.docx"
    document = WordDocument()
    table = document.add_table(rows=3, cols=3)
    table.cell(0, 0).merge(table.cell(0, 1)).text = "Merged header"
    table.cell(0, 2).text = "Value"
    table.cell(1, 0).merge(table.cell(2, 0)).text = "Vertical"
    table.cell(1, 1).text = "r1c2"
    table.cell(1, 2).text = "r1c3"
    table.cell(2, 1).text = "r2c2"
    table.cell(2, 2).text = "r2c3"
    document.save(source)

    parsed = DocumentParser().parse_structured(source, document_id="doc-word", version=1)

    assert [element.text for element in parsed.elements] == [
        "| Merged header |  | Value |\n| --- | --- | --- |\n| Vertical | r1c2 | r1c3 |",
        "| Merged header |  | Value |\n| --- | --- | --- |\n|  | r2c2 | r2c3 |",
    ]


def test_pdf_parser_keeps_page_block_order_and_removes_repeated_margins(tmp_path):
    import fitz

    source = tmp_path / "guide.pdf"
    pdf = fitz.open()
    for page_number in range(1, 4):
        page = pdf.new_page(width=595, height=842)
        page.insert_text((60, 30), "Product Guide", fontsize=9)
        page.insert_text((60, 110), f"Section {page_number}", fontsize=18)
        page.insert_text((60, 150), f"Body paragraph {page_number}.", fontsize=11)
        page.insert_text((270, 820), f"Page {page_number}", fontsize=9)
    pdf.save(source)
    pdf.close()

    parsed = DocumentParser().parse_structured(source, document_id="doc-pdf", version=1)

    assert [element.text for element in parsed.elements] == [
        "Section 1",
        "Body paragraph 1.",
        "Section 2",
        "Body paragraph 2.",
        "Section 3",
        "Body paragraph 3.",
    ]
    assert [element.page for element in parsed.elements] == [1, 1, 2, 2, 3, 3]
    assert parsed.elements[0].element_type == "heading"
    assert parsed.elements[1].section_path == ("Section 1",)
    assert parsed.elements[0].bounding_box == pytest.approx(
        {"x0": 60.0, "y0": 90.65, "x1": 135.04, "y1": 115.38}, abs=0.1
    )


def test_pdf_parser_joins_visual_line_wraps_without_crossing_block_or_list_boundaries(
    monkeypatch,
):
    import fitz

    def line(text, y0, y1=12):
        return {
            "bbox": (20, y0, 300, y0 + y1),
            "spans": [{"text": text, "size": 11, "font": "Regular", "flags": 0}],
        }

    page = SimpleNamespace(
        rect=SimpleNamespace(height=800),
        get_text=lambda _format: {
            "blocks": [
                {
                    "type": 0,
                    "bbox": (20, 100, 300, 130),
                    "lines": [line("这是中文", 100), line("视觉换行。", 113)],
                },
                {
                    "type": 0,
                    "bbox": (20, 150, 300, 180),
                    "lines": [line("English visual", 150), line("line wrap.", 163)],
                },
                {
                    "type": 0,
                    "bbox": (20, 200, 300, 230),
                    "lines": [line("hyphen-", 200), line("ated word.", 213)],
                },
                {
                    "type": 0,
                    "bbox": (20, 250, 300, 280),
                    "lines": [line("- first", 250), line("- second", 263)],
                },
                {
                    "type": 0,
                    "bbox": (20, 300, 300, 355),
                    "lines": [line("Paragraph one.", 300), line("Paragraph two.", 343)],
                },
            ]
        },
    )

    class FakePdf:
        def __iter__(self):
            return iter([page])

        def close(self):
            pass

    monkeypatch.setattr(fitz, "open", lambda *args, **kwargs: FakePdf())

    from KBzhy.app.core.parsers.pdf_parser import parse_pdf

    elements = parse_pdf(b"pdf", document_id="doc-pdf")

    assert [element.text for element in elements] == [
        "这是中文视觉换行。",
        "English visual line wrap.",
        "hyphenated word.",
        "- first\n- second",
        "Paragraph one.\nParagraph two.",
    ]


def test_pdf_parser_skips_dirty_items_without_losing_valid_blocks(monkeypatch):
    import fitz

    page = SimpleNamespace(
        rect=SimpleNamespace(height=800),
        get_text=lambda _format: {
            "blocks": [
                {
                    "type": 0,
                    "bbox": (20, 100, 300, 130),
                    "lines": [
                        {
                            "bbox": (20, 100, 300, 112),
                            "spans": [
                                {"text": "Valid ", "size": 11, "font": "Regular", "flags": 0},
                                {"text": "dirty", "size": float("nan"), "font": "Bad", "flags": object()},
                                {"text": "text", "size": 11, "font": "Regular", "flags": 0},
                            ],
                        },
                        {"bbox": (1, 2), "spans": [{"text": "bad line", "size": 11}]},
                    ],
                },
                {
                    "type": 0,
                    "bbox": (20, float("inf"), 300, 180),
                    "lines": [
                        {
                            "bbox": (20, 150, 300, 162),
                            "spans": [{"text": "Second block.", "size": 11, "font": "Regular", "flags": 0}],
                        }
                    ],
                },
            ]
        },
    )

    class FakePdf:
        def __iter__(self):
            return iter([page])

        def close(self):
            pass

    monkeypatch.setattr(fitz, "open", lambda *args, **kwargs: FakePdf())

    from KBzhy.app.core.parsers.pdf_parser import parse_pdf

    elements = parse_pdf(b"pdf", document_id="doc-pdf")

    assert [element.text for element in elements] == ["Valid text", "Second block."]
    assert all(
        all(math.isfinite(value) for value in element.bounding_box.values())
        for element in elements
    )


def test_pdf_heading_stack_uses_comparable_levels(monkeypatch):
    import fitz

    def block(text, size, y):
        return {
            "type": 0,
            "bbox": (20, y, 300, y + 20),
            "lines": [
                {
                    "bbox": (20, y, 300, y + 20),
                    "spans": [{"text": text, "size": size, "font": "Bold", "flags": 16}],
                }
            ],
        }

    page = SimpleNamespace(
        rect=SimpleNamespace(height=800),
        get_text=lambda _format: {
            "blocks": [
                block("Deep", 14, 100),
                block("Parent", 18, 150),
                block("Body paragraph.", 11, 200),
            ]
        },
    )

    class FakePdf:
        def __iter__(self):
            return iter([page])

        def close(self):
            pass

    monkeypatch.setattr(fitz, "open", lambda *args, **kwargs: FakePdf())

    from KBzhy.app.core.parsers.pdf_parser import parse_pdf

    elements = parse_pdf(b"pdf", document_id="doc-pdf")

    assert elements[0].metadata["level"] > elements[1].metadata["level"]
    assert elements[1].section_path == ("Parent",)
    assert elements[2].section_path == ("Parent",)


@pytest.mark.parametrize("failure_stage", ["iteration", "get_text"])
def test_pdf_parser_wraps_processing_errors_and_closes_document(monkeypatch, failure_stage):
    import fitz

    closed = False

    class FailingPage:
        rect = SimpleNamespace(height=800)

        def get_text(self, _format):
            if failure_stage == "get_text":
                raise RuntimeError("get_text failed")
            return {
                "blocks": [
                    {
                        "type": 0,
                        "bbox": (object(), 1, 2, 3),
                        "lines": [
                            {
                                "bbox": (1, 1, 2, 2),
                                "spans": [
                                    {"text": "body", "size": 11, "font": "Regular", "flags": 0}
                                ],
                            }
                        ],
                    }
                ]
            }

    class FakePdf:
        def __iter__(self):
            if failure_stage == "iteration":
                raise RuntimeError("page iteration failed")
            return iter([FailingPage()])

        def close(self):
            nonlocal closed
            closed = True

    monkeypatch.setattr(fitz, "open", lambda *args, **kwargs: FakePdf())

    from KBzhy.app.core.parsers.pdf_parser import parse_pdf

    with pytest.raises(DocumentParseError, match="PDF 文档解析失败") as caught:
        parse_pdf(b"pdf", document_id="doc-pdf")

    assert isinstance(caught.value.__cause__, (RuntimeError, TypeError))
    assert closed is True


@pytest.mark.parametrize("failure_stage", ["body_iteration", "table_extraction"])
def test_word_parser_wraps_processing_errors(monkeypatch, failure_stage):
    import docx
    import docx.table

    class FailingBody:
        def iterchildren(self):
            if failure_stage == "body_iteration":
                raise RuntimeError("body iteration failed")
            return [SimpleNamespace(tag="{word}tbl")]

    fake_document = SimpleNamespace(element=SimpleNamespace(body=FailingBody()))
    monkeypatch.setattr(docx, "Document", lambda _source: fake_document)

    if failure_stage == "table_extraction":
        monkeypatch.setattr(
            docx.table,
            "Table",
            lambda *_args: (_ for _ in ()).throw(RuntimeError("table extraction failed")),
        )

    from KBzhy.app.core.parsers.word_parser import parse_word

    with pytest.raises(DocumentParseError, match="Word 文档解析失败") as caught:
        parse_word(b"docx", document_id="doc-word")

    assert isinstance(caught.value.__cause__, RuntimeError)


def test_scanned_pdf_is_rejected_instead_of_indexing_placeholder(tmp_path):
    import fitz

    source = tmp_path / "scan.pdf"
    pdf = fitz.open()
    pdf.new_page()
    pdf.save(source)
    pdf.close()

    with pytest.raises(DocumentParseError, match="未提取到可索引文字"):
        DocumentParser().parse_structured(source, document_id="doc-scan", version=1)


def test_legacy_documents_become_ordered_paragraph_and_table_elements(monkeypatch):
    parser = DocumentParser()
    legacy_documents = [
        Document("first paragraph", {"page": 3, "source": "guide.pdf"}),
        Document("name | value", {"sheet": "Sheet1", "source": "guide.xlsx"}),
    ]
    monkeypatch.setattr(parser, "parse", lambda _source: legacy_documents)

    parsed = parser.parse_structured(
        "ignored.txt",
        document_id="doc1",
        version=1,
        kb_id="kb1",
    )

    assert [element.element_type for element in parsed.elements] == ["paragraph", "table"]
    assert [element.order for element in parsed.elements] == [0, 1]
    assert parsed.elements[0].page == 3
    assert parsed.elements[1].metadata["sheet"] == "Sheet1"


def test_parsed_artifact_round_trip_is_lossless(tmp_path):
    parser = DocumentParser(artifact_dir=tmp_path)
    parsed = ParsedDocument(
        document_id="doc1",
        version=4,
        title="指南",
        language="zh-CN",
        elements=(
            DocumentElement(
                element_id="doc1:0",
                element_type="paragraph",
                text="正文",
                order=0,
                page=2,
                section_path=("第一章", "概述"),
                bounding_box={"x0": 1.25, "y0": 2.5},
                metadata={"bold": True, "nested": {"values": [1, 2]}},
            ),
        ),
        metadata={"kb_id": "kb1", "source": "指南.md", "tags": ["a", "b"]},
    )

    path = parser.save_artifact(parsed)

    assert path == tmp_path / "kb1" / "doc1" / "v4.json"
    assert parser.load_artifact(path) == parsed


def test_load_artifact_rejects_path_outside_artifact_directory(tmp_path):
    parser = DocumentParser(artifact_dir=tmp_path / "artifacts")
    outside = tmp_path / "outside.json"
    outside.write_text("{}", encoding="utf-8")

    with pytest.raises(ValueError, match="artifact_dir"):
        parser.load_artifact(outside)


@pytest.mark.parametrize(
    "payload",
    [
        "not-json",
        "[]",
        '{"document_id":"doc1","version":1,"elements":[],"metadata":{}}',
        '{"document_id":1,"version":1,"elements":[],"metadata":{}}',
        '{"document_id":"doc1","version":"1","elements":[],"metadata":{}}',
        '{"document_id":"doc1","version":1,"elements":[],"metadata":[]}',
        '{"document_id":"doc1","version":1,"elements":{},"metadata":{}}',
        (
            '{"document_id":"doc1","version":1,"title":"guide",'
            '"language":"und","elements":[{"element_id":"e1",'
            '"element_type":"image","text":"body","order":0}],"metadata":{}}'
        ),
        (
            '{"document_id":"doc1","version":1,"title":"guide",'
            '"language":"und","elements":["paragraph"],"metadata":{}}'
        ),
        (
            '{"document_id":"doc1","version":1,"title":"guide",'
            '"language":"und","elements":[{"element_id":"e1",'
            '"element_type":"paragraph","text":"body","order":"0"}],"metadata":{}}'
        ),
    ],
)
def test_load_artifact_rejects_malformed_artifacts_with_stable_error(tmp_path, payload):
    artifact_dir = tmp_path / "artifacts"
    artifact_dir.mkdir()
    artifact = artifact_dir / "bad.json"
    artifact.write_text(payload, encoding="utf-8")
    parser = DocumentParser(artifact_dir=artifact_dir)

    with pytest.raises(ParsedArtifactError, match="invalid parsed artifact"):
        parser.load_artifact(artifact)


def test_save_artifact_uses_atomic_replace_and_cleans_temp_on_failure(tmp_path, monkeypatch):
    parser = DocumentParser(artifact_dir=tmp_path)
    parsed = ParsedDocument(
        document_id="doc1",
        version=1,
        title="guide",
        language="und",
        metadata={"kb_id": "kb1"},
    )

    def fail_replace(_source, _target):
        raise OSError("replace failed")

    monkeypatch.setattr("KBzhy.app.core.parser.os.replace", fail_replace)

    with pytest.raises(OSError, match="replace failed"):
        parser.save_artifact(parsed)

    target_dir = tmp_path / "kb1" / "doc1"
    assert not (target_dir / "v1.json").exists()
    assert list(target_dir.glob("*.tmp")) == []


def test_save_artifact_cleans_temp_when_serialization_fails(tmp_path):
    parser = DocumentParser(artifact_dir=tmp_path)
    parsed = ParsedDocument(
        document_id="doc1",
        version=1,
        title="guide",
        language="und",
        metadata={"kb_id": "kb1", "unsupported": object()},
    )

    with pytest.raises(TypeError):
        parser.save_artifact(parsed)

    target_dir = tmp_path / "kb1" / "doc1"
    assert list(target_dir.glob("*.tmp")) == []


def test_save_artifact_rejects_symlink_escape(tmp_path):
    artifact_dir = tmp_path / "artifacts"
    outside = tmp_path / "outside"
    artifact_dir.mkdir()
    outside.mkdir()
    link = artifact_dir / "kb1"
    junction_created = False
    try:
        try:
            link.symlink_to(outside, target_is_directory=True)
        except OSError as exc:
            if os.name != "nt":
                pytest.skip(f"directory symlinks unavailable: {exc}")
            result = subprocess.run(
                ["cmd.exe", "/d", "/c", "mklink", "/J", str(link), str(outside)],
                capture_output=True,
                text=True,
                check=False,
            )
            if result.returncode != 0:
                pytest.skip(f"directory links unavailable: {result.stderr}")
            junction_created = True

        parser = DocumentParser(artifact_dir=artifact_dir)
        parsed = ParsedDocument(
            document_id="doc1",
            version=1,
            title="guide",
            language="und",
            metadata={"kb_id": "kb1"},
        )

        with pytest.raises(ValueError, match="artifact_dir"):
            parser.save_artifact(parsed)

        assert not (outside / "doc1" / "v1.json").exists()
    finally:
        if junction_created and link.exists():
            os.rmdir(link)


@pytest.mark.parametrize("field,value", [("kb_id", "../kb"), ("document_id", "doc/child")])
def test_artifact_path_rejects_unsafe_identifiers(tmp_path, field, value):
    parser = DocumentParser(artifact_dir=tmp_path)
    values = {"kb_id": "kb1", "document_id": "doc1"}
    values[field] = value
    parsed = ParsedDocument(
        document_id=values["document_id"],
        version=1,
        title="guide",
        language="und",
        metadata={"kb_id": values["kb_id"]},
    )

    with pytest.raises(ValueError, match=field):
        parser.save_artifact(parsed)


@pytest.mark.parametrize("version", [0, -1, True])
def test_parse_structured_rejects_invalid_version(tmp_path, version):
    parser = DocumentParser()

    with pytest.raises(ValueError, match="version"):
        parser.parse_structured(
            _markdown_fixture(tmp_path),
            document_id="doc1",
            version=version,
            kb_id="kb1",
        )


@pytest.mark.parametrize(
    "raw_section_path,expected",
    [(None, ()), ("Overview", ("Overview",)), (["A", "B"], ("A", "B")), (("A",), ("A",))],
)
def test_legacy_section_path_is_normalized(monkeypatch, raw_section_path, expected):
    parser = DocumentParser()
    monkeypatch.setattr(
        parser,
        "parse",
        lambda _source: [Document("body", {"section_path": raw_section_path})],
    )

    parsed = parser.parse_structured("ignored.txt", document_id="doc1", version=1)

    assert parsed.elements[0].section_path == expected


def test_legacy_section_path_rejects_unsupported_type(monkeypatch):
    parser = DocumentParser()
    monkeypatch.setattr(
        parser,
        "parse",
        lambda _source: [Document("body", {"section_path": {"heading": "A"}})],
    )

    with pytest.raises(ParsedArtifactError, match="section_path"):
        parser.parse_structured("ignored.txt", document_id="doc1", version=1)
